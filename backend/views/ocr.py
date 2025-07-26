from django.http import JsonResponse
from rest_framework.views import APIView
from backend.models.file_conversion import DocumentConversion, SupportedConversion
from backend.models.user import CustomToken
from backend.ocr_extractors import OCRProcessor
from django.contrib.auth.models import AnonymousUser
from django.core.files.base import ContentFile
from django.utils import timezone
import io


class OCRView(APIView):
    def __init__(self):
        super().__init__()
        self.ocr_processor = OCRProcessor()
    
    def post(self, request):
        auth_token = request.POST.get("auth_token")

        if auth_token == 'null':
            user = None
        else:
            try:
                token = CustomToken.objects.get(key=auth_token)
                if token.is_expired():
                    token.delete()
                    user = None
                else:
                    user = token.user

            except CustomToken.DoesNotExist:
                user = None

        original_file = request.FILES.get("original_file", None)
        if not original_file:
            return JsonResponse({"error": "original_file is required"}, status=400)

        converted_mimetype = request.POST.get("converted_mimetype")
        if not converted_mimetype:
            return JsonResponse({"error": "converted_mimetype is required"}, status=400)

        original_mimetype = SupportedConversion.get_original_mimetype(
            original_file.content_type, original_file.name.split('.')[-1])

        # Validate supported version for OCR
        ocr_enabled_conversions = SupportedConversion.ocr_enabled_conversions()
        supported_conversion = ocr_enabled_conversions.filter(
            original_mimetype=original_mimetype,
            target_mimetype=converted_mimetype,
        )
        
        if not supported_conversion.exists():
            return JsonResponse({
                "error": f"OCR conversion from {original_mimetype} to {converted_mimetype} is not supported"
            }, status=400)

        user = request.user if not isinstance(
            request.user, AnonymousUser) else None

        document_conversion = DocumentConversion.objects.create(
            original_file=original_file,
            original_filename=original_file.name,
            original_mimetype=original_mimetype,
            original_size=original_file.size,
            converted_mimetype=converted_mimetype,
            status="processing",
            user=user,
        )

        try:
            # Extract text from the input file using OCR processor
            extracted_text = self.ocr_processor.extract_text(original_file, original_mimetype)
            
            if not extracted_text:
                document_conversion.status = 'failed'
                document_conversion.error_message = 'No text could be extracted from the file'
                document_conversion.save()
                return JsonResponse({'error': 'No text could be extracted from the file'}, status=500)

            # Create new file with extracted text in the converted mimetype
            converted_file = self.embed_text_in_new_file(extracted_text, converted_mimetype, original_file.name)
            
            if not converted_file:
                document_conversion.status = 'failed'
                document_conversion.error_message = 'Failed to create file with extracted text'
                document_conversion.save()
                return JsonResponse({'error': 'Failed to create file with extracted text'}, status=500)

            # Get the target extension for the converted file
            supported_conversion_obj = supported_conversion.first()
            converted_filename = f"{original_file.name.rsplit('.')[0]}.{supported_conversion_obj.target_extension}"

            # Save the converted file
            document_conversion.converted_file.save(
                converted_filename, converted_file, save=False)

            # Update model fields
            document_conversion.converted_filename = document_conversion.converted_file.name.split('/')[-1]
            document_conversion.converted_size = document_conversion.converted_file.size
            document_conversion.completed_at = timezone.now()
            document_conversion.status = 'completed'
            document_conversion.save()

            response_data = document_conversion.to_dict()
            response_data['extracted_text'] = extracted_text
            
            return JsonResponse(response_data, status=200)

        except Exception as e:
            document_conversion.status = 'failed'
            document_conversion.error_message = str(e)
            document_conversion.save()
            return JsonResponse({'error': f'OCR processing failed: {str(e)}'}, status=500)

    def get(self, request):
        """
        Get information about supported OCR extractors and MIME types.
        """
        try:
            info = {
                'supported_mimetypes': self.ocr_processor.get_supported_mimetypes(),
                'extractors': self.ocr_processor.get_extractor_info(),
                'total_extractors': len(self.ocr_processor.extractors)
            }
            return JsonResponse(info, status=200)
        except Exception as e:
            return JsonResponse({'error': f'Failed to get OCR info: {str(e)}'}, status=500)


    def embed_text_in_new_file(self, text, target_mimetype, original_filename):
        """Create a new file with the extracted text in the specified format"""
        try:
            if target_mimetype == 'text/plain':
                return ContentFile(text.encode('utf-8'))
            
            elif target_mimetype == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Create DOCX file
                from docx import Document
                doc = Document()
                doc.add_paragraph(text)
                
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                return ContentFile(doc_buffer.getvalue())
            
            elif target_mimetype == 'application/msword':
                # Create DOC file - similar to DOCX for now
                from docx import Document
                doc = Document()
                doc.add_paragraph(text)
                
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                return ContentFile(doc_buffer.getvalue())
            
            elif target_mimetype == 'application/pdf':
                # Create PDF file
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                from reportlab.lib.utils import simpleSplit
                
                pdf_buffer = io.BytesIO()
                c = canvas.Canvas(pdf_buffer, pagesize=letter)
                width, height = letter
                
                # Split text into lines that fit the page
                lines = text.split('\n')
                y_position = height - 50
                
                for line in lines:
                    if y_position < 50:  # Start new page if needed
                        c.showPage()
                        y_position = height - 50
                    
                    # Split long lines
                    wrapped_lines = simpleSplit(line, 'Helvetica', 12, width - 100)
                    for wrapped_line in wrapped_lines:
                        if y_position < 50:
                            c.showPage()
                            y_position = height - 50
                        c.drawString(50, y_position, wrapped_line)
                        y_position -= 15
                
                c.save()
                pdf_buffer.seek(0)
                
                return ContentFile(pdf_buffer.getvalue())
            
            else:
                # Default to plain text
                return ContentFile(text.encode('utf-8'))
                
        except Exception as e:
            print(f"Error creating file with extracted text: {e}")
            return None